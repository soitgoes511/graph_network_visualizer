import io
import re
import xml.etree.ElementTree as ET

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

try:
    import openpyxl
except Exception:  # pragma: no cover - dependency availability is environment-specific
    openpyxl = None

try:
    import xlrd
except Exception:  # pragma: no cover - dependency availability is environment-specific
    xlrd = None

URL_PATTERN = r"http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+"
MAX_TEXT_CHARS = 250000
MAX_EXCEL_SHEETS = 12
MAX_EXCEL_ROWS = 5000
MAX_EXCEL_COLS = 48
MAX_XML_NODES = 25000


def extract_links(text):
    return re.findall(URL_PATTERN, text or "")


def _dedupe_links(links):
    seen = set()
    unique = []
    for link in links:
        item = str(link or "").strip()
        if not item:
            continue
        if item in seen:
            continue
        seen.add(item)
        unique.append(item)
    return unique


def _compact_text(text: str, limit: int = MAX_TEXT_CHARS) -> str:
    compact = re.sub(r"\s+", " ", text or "").strip()
    if len(compact) <= limit:
        return compact
    return compact[:limit]


def parse_pdf(file_content: bytes, logger=None):
    if logger:
        logger("Parsing PDF document...")
    text = ""
    links = []
    try:
        reader = PdfReader(io.BytesIO(file_content))
        total_pages = len(reader.pages)
        if logger:
            logger(f"PDF has {total_pages} pages")

        for i, page in enumerate(reader.pages):
            if logger and i % 5 == 0:
                logger(f"Processing page {i + 1}/{total_pages}")

            text += (page.extract_text() or "") + "\n"
            if "/Annots" in page:
                for annot in page["/Annots"]:
                    obj = annot.get_object()
                    if "/A" in obj and "/URI" in obj["/A"]:
                        links.append(obj["/A"]["/URI"])
    except Exception as error:
        if logger:
            logger(f"Error parsing PDF: {error}")
        print(f"Error parsing PDF: {error}")

    links.extend(extract_links(text))
    text = _compact_text(text)

    if logger:
        logger(f"PDF parsing complete. Found {len(links)} links.")
    return {"text": text, "links": _dedupe_links(links)}


def parse_docx(file_content: bytes, logger=None):
    if logger:
        logger("Parsing DOCX document...")
    text = ""
    links = []
    try:
        doc = Document(io.BytesIO(file_content))
        if logger:
            logger(f"DOCX has {len(doc.paragraphs)} paragraphs")

        for para in doc.paragraphs:
            text += para.text + "\n"

        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                links.append(rel.target_ref)

        links.extend(extract_links(text))

    except Exception as error:
        if logger:
            logger(f"Error parsing DOCX: {error}")
        print(f"Error parsing DOCX: {error}")

    text = _compact_text(text)
    if logger:
        logger(f"DOCX parsing complete. Found {len(links)} links.")
    return {"text": text, "links": _dedupe_links(links)}


def parse_excel(file_content: bytes, filename: str, logger=None):
    name = str(filename or "").lower()
    if logger:
        logger(f"Parsing Excel document ({name or 'unknown'})...")

    rows_text = []
    links = []

    if name.endswith(".xlsx"):
        if openpyxl is None:
            raise RuntimeError("openpyxl is required to parse .xlsx files.")

        workbook = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
        for sheet_index, sheet_name in enumerate(workbook.sheetnames[:MAX_EXCEL_SHEETS]):
            worksheet = workbook[sheet_name]
            rows_text.append(f"SHEET: {sheet_name}")
            for row_index, row in enumerate(worksheet.iter_rows(min_row=1, max_row=MAX_EXCEL_ROWS, max_col=MAX_EXCEL_COLS), start=1):
                values = []
                for cell in row:
                    cell_value = cell.value
                    if cell_value is None:
                        continue
                    value_text = str(cell_value).strip()
                    if value_text:
                        values.append(value_text)
                        links.extend(extract_links(value_text))
                    if cell.hyperlink and cell.hyperlink.target:
                        links.append(str(cell.hyperlink.target))

                if values:
                    rows_text.append(" | ".join(values))

                if logger and row_index % 1000 == 0:
                    logger(f"Processed {row_index} rows from sheet {sheet_index + 1}/{min(len(workbook.sheetnames), MAX_EXCEL_SHEETS)}")

        workbook.close()

    elif name.endswith(".xls"):
        if xlrd is None:
            raise RuntimeError("xlrd is required to parse .xls files.")

        workbook = xlrd.open_workbook(file_contents=file_content)
        sheet_count = min(workbook.nsheets, MAX_EXCEL_SHEETS)
        for sheet_index in range(sheet_count):
            worksheet = workbook.sheet_by_index(sheet_index)
            rows_text.append(f"SHEET: {worksheet.name}")
            max_rows = min(worksheet.nrows, MAX_EXCEL_ROWS)
            max_cols = min(worksheet.ncols, MAX_EXCEL_COLS)

            for row_index in range(max_rows):
                values = []
                for col_index in range(max_cols):
                    cell_value = worksheet.cell_value(row_index, col_index)
                    if cell_value in {"", None}:
                        continue

                    value_text = str(cell_value).strip()
                    if value_text:
                        values.append(value_text)
                        links.extend(extract_links(value_text))

                if values:
                    rows_text.append(" | ".join(values))

                if logger and row_index % 1000 == 0 and row_index > 0:
                    logger(f"Processed {row_index} rows from sheet {sheet_index + 1}/{sheet_count}")

    else:
        raise RuntimeError("Unsupported Excel extension. Expected .xlsx or .xls")

    text = _compact_text("\n".join(rows_text))
    links.extend(extract_links(text))
    if logger:
        logger(f"Excel parsing complete. Found {len(links)} links.")
    return {"text": text, "links": _dedupe_links(links)}


def parse_html(file_content: bytes, logger=None):
    if logger:
        logger("Parsing HTML document...")

    decoded = file_content.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(decoded, "html.parser")

    for tag in soup.find_all(["script", "style", "noscript"]):
        tag.decompose()

    links = []
    for anchor in soup.find_all("a", href=True):
        href = str(anchor.get("href", "")).strip()
        if href.startswith("http"):
            links.append(href)

    body = soup.body or soup
    text_blocks = []
    for element in body.find_all(["title", "h1", "h2", "h3", "p", "li", "td"], limit=12000):
        text = element.get_text(" ", strip=True)
        if text:
            text_blocks.append(text)

    text = _compact_text("\n".join(text_blocks))
    links.extend(extract_links(text))

    if logger:
        logger(f"HTML parsing complete. Found {len(links)} links.")
    return {"text": text, "links": _dedupe_links(links)}


def parse_xml(file_content: bytes, logger=None):
    if logger:
        logger("Parsing XML document...")

    decoded = file_content.decode("utf-8", errors="ignore")
    links = []
    text_chunks = []

    try:
        root = ET.fromstring(decoded)
        for index, elem in enumerate(root.iter()):
            if index >= MAX_XML_NODES:
                break

            if elem.text and elem.text.strip():
                text_chunks.append(elem.text.strip())
                links.extend(extract_links(elem.text))

            for attr_value in elem.attrib.values():
                attr_text = str(attr_value).strip()
                if attr_text:
                    text_chunks.append(attr_text)
                    links.extend(extract_links(attr_text))
                    if attr_text.startswith("http"):
                        links.append(attr_text)
    except Exception:
        soup = BeautifulSoup(decoded, "xml")
        text_chunks.append(soup.get_text(" ", strip=True))
        for tag in soup.find_all(href=True):
            href = str(tag.get("href", "")).strip()
            if href:
                links.append(href)

    text = _compact_text("\n".join(text_chunks))
    links.extend(extract_links(text))

    if logger:
        logger(f"XML parsing complete. Found {len(links)} links.")
    return {"text": text, "links": _dedupe_links(links)}
